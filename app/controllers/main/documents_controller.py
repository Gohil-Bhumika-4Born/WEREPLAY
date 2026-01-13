"""
Documents controller for handling documents-related requests.
"""
import os
from flask import render_template, request, redirect, url_for, flash, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename

from app.extensions import db
from app.models.documents import Document
import PyPDF2
import docx

# ==============================
# CHANGED HERE: unified extractor
# ==============================
def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    try:
        # -------- PDF --------
        if ext == '.pdf':
            text = ''
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ''
            return text.strip()

        # -------- TXT --------
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()

        # -------- DOCX --------
        if ext == '.docx':
            doc = docx.Document(file_path)
            return '\n'.join(p.text for p in doc.paragraphs).strip()

    except Exception as e:
        # CHANGED HERE: fail safely, never crash upload
        print('Text extraction failed:', e)

    return None


class DocumentsController:

    @staticmethod
    @login_required
    def documents_upload_page():
        """
        Render upload page (GET only)
        """
        return render_template('main/documents-upload.html')

    @staticmethod
    @login_required
    def upload_document():
        """
        Handle document upload (POST)
        """

        # ==============================
        # CHANGED HERE: correct field name
        # ==============================
        file = request.files.get('documents')

        if not file or file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)

        filename = secure_filename(file.filename)
        file_ext = filename.rsplit('.', 1)[-1].upper()

        upload_dir = os.path.join(
            current_app.root_path, 'static', 'uploads', 'documents'
        )
        os.makedirs(upload_dir, exist_ok=True)

        file_path = os.path.join(upload_dir, filename)
        file.save(file_path)

        # ==============================
        # CHANGED HERE: real file size
        # ==============================
        file_size = os.path.getsize(file_path)

        # ==============================
        # CHANGED HERE: auto title fallback
        # ==============================
        title = request.form.get('title') or filename.rsplit('.', 1)[0]

        # ==============================
        # CHANGED HERE: REAL EXTRACTION
        # ==============================
        extracted_text = extract_text_from_file(file_path)

        # ==============================
        # CHANGED HERE: description fallback
        # ==============================
        description = request.form.get('description')
        if not description and extracted_text:
            description = extracted_text[:300]

        # ==============================
        # CHANGED HERE: save full data
        # ==============================
        document = Document(
            document_title=title,
            category=request.form.get('category'),
            description=description,
            tags=request.form.get('tags'),
            document_path=f'uploads/documents/{filename}',
            file_type=file_ext,
            file_size=file_size,
            extracted_text=extracted_text,   # 🔥 STEP-1 CORE
            user_id=current_user.id
        )

        db.session.add(document)
        db.session.commit()

        flash('Document uploaded and processed', 'success')
        return redirect(url_for('main.documents'))

    @staticmethod
    @login_required
    def documents_page():
        """
        Documents list page
        """
        documents = Document.query.filter_by(
            user_id=current_user.id
        ).order_by(Document.created_at.desc()).all()

        return render_template(
            'main/documents.html',
            documents=documents
        )
